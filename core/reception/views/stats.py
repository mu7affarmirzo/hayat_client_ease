from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.db.models import Sum, Count, F, ExpressionWrapper, DecimalField
from django.http import HttpResponse
from datetime import datetime, timedelta
import csv
import io
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from core.models import SessionModel, Account, TherapistModel, PaymentModel

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


@login_required
def therapist_statistics(request):
    """
    View to display therapist statistics for a specified date range
    with new profit distribution logic based on actual paid amounts
    """
    # Default date range (last 30 days)
    end_date = datetime.now().date()
    start_date = end_date - timedelta(days=30)
    date_range = f"{start_date.strftime('%d/%m/%Y')} - {end_date.strftime('%d/%m/%Y')}"

    # Get date range from request if provided
    if 'date_range' in request.GET and request.GET['date_range']:
        date_range = request.GET['date_range']
        start_str, end_str = date_range.split(' - ')
        start_date = datetime.strptime(start_str, '%d/%m/%Y').date()
        end_date = datetime.strptime(end_str, '%d/%m/%Y').date()

    # Get all therapists
    therapists = TherapistModel.objects.all()

    # Process statistics for each therapist
    stats = []
    total_sessions = 0
    total_amount = 0
    total_payout = 0

    for therapist in therapists:
        # Get sessions for this therapist in the date range
        sessions = SessionModel.objects.filter(
            therapist=therapist,
            created_at__date__gte=start_date,
            created_at__date__lte=end_date,
        )

        # Calculate total paid amount for these sessions
        session_ids = sessions.values_list('id', flat=True)
        paid_amount = PaymentModel.objects.filter(
            session_id__in=session_ids
        ).aggregate(total=Sum('amount'))['total'] or 0

        # Calculate statistics
        session_count = sessions.aggregate(total=Sum('proceeded_sessions'))['total'] or 0
        # Skip therapists with no sessions or no payments
        if paid_amount == 0:
            continue

        # Add to totals
        total_sessions += session_count
        total_amount += paid_amount

        # Calculate payout based on therapist rate (from one owner's share)
        rate = therapist.rate if hasattr(therapist, 'rate') else 0

        # New calculation: therapist gets their rate % from HALF of the paid amount
        owner_share = paid_amount // 2
        payout = owner_share * (rate / 100)
        total_payout += payout

        # Add to stats list
        stats.append({
            'full_name': therapist.full_name,
            'total_amount': total_amount,
            'rate': rate,
            'session_count': session_count,
            'after_share_amount': owner_share,
            'payout_amount': payout,
        })

    # Calculate split between owners
    owner1_share = total_amount // 2
    owner2_share = total_amount // 2

    # Owner 1 pays the therapists
    owner1_final = int(owner1_share) - total_payout

    context = {
        'therapists': stats,
        'start_date': start_date.strftime('%d/%m/%Y'),
        'end_date': end_date.strftime('%d/%m/%Y'),
        'date_range': date_range,
        'total_sessions': total_sessions,
        'total_amount': total_amount,
        'total_payout': total_payout,
        'owner1_share': owner1_share,
        'owner2_share': owner2_share,
        'owner1_final': owner1_final,
    }

    return render(request, 'adminstration/therapist_stats.html', context)


@login_required
def export_therapist_statistics(request):
    """
    Export therapist statistics to Excel with new profit distribution logic
    based on actual paid amounts
    """
    # Default date range (last 30 days)
    end_date = datetime.now().date()
    start_date = end_date - timedelta(days=30)

    # Get date range from request if provided
    if 'date_range' in request.GET and request.GET['date_range']:
        date_range = request.GET['date_range']
        start_str, end_str = date_range.split(' - ')
        start_date = datetime.strptime(start_str, '%d/%m/%Y').date()
        end_date = datetime.strptime(end_str, '%d/%m/%Y').date()

    # Get all therapists
    therapists = TherapistModel.objects.all()

    # Process statistics for each therapist
    stats = []
    total_sessions = 0
    total_amount = 0
    total_payout = 0

    for therapist in therapists:
        # Get sessions for this therapist in the date range
        sessions = SessionModel.objects.filter(
            therapist=therapist,
            created_at__date__gte=start_date,
            created_at__date__lte=end_date,
        )

        # Calculate total paid amount for these sessions
        session_ids = sessions.values_list('id', flat=True)
        paid_amount = PaymentModel.objects.filter(
            session_id__in=session_ids
        ).aggregate(total=Sum('amount'))['total'] or 0

        # Calculate statistics
        session_count = sessions.aggregate(count=Count('id'))['count'] or 0

        # Skip therapists with no sessions or no payments
        if session_count == 0 or paid_amount == 0:
            continue

        # Add to totals
        total_sessions += session_count
        total_amount += paid_amount

        # Calculate payout based on therapist rate (from one owner's share)
        rate = therapist.rate if hasattr(therapist, 'rate') else 0

        # New calculation: therapist gets their rate % from HALF of the paid amount
        owner_share = paid_amount / 2
        payout = owner_share * (rate / 100)
        total_payout += payout

        # Add to stats list
        stats.append({
            'full_name': therapist.full_name,
            'rate': rate,
            'session_count': session_count,
            'total_amount': paid_amount,
            'payout_amount': payout,
        })

    # Calculate split between owners
    owner1_share = total_amount / 2
    owner2_share = total_amount / 2

    # Owner 1 pays the therapists
    owner1_final = owner1_share - total_payout

    # Create Excel workbook
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Статистика терапевтов"

    # Set column widths
    worksheet.column_dimensions['A'].width = 5
    worksheet.column_dimensions['B'].width = 25
    worksheet.column_dimensions['C'].width = 15
    worksheet.column_dimensions['D'].width = 15
    worksheet.column_dimensions['E'].width = 20
    worksheet.column_dimensions['F'].width = 20

    # Add title
    title = f"Статистика терапевтов за период {start_date.strftime('%d/%m/%Y')} - {end_date.strftime('%d/%m/%Y')}"
    worksheet['A1'] = title
    worksheet.merge_cells('A1:F1')
    title_cell = worksheet['A1']
    title_cell.font = Font(size=14, bold=True)
    title_cell.alignment = Alignment(horizontal='center')

    # Add subtitle about calculation method
    subtitle = "Расчет на основе фактически оплаченных сумм"
    worksheet['A2'] = subtitle
    worksheet.merge_cells('A2:F2')
    subtitle_cell = worksheet['A2']
    subtitle_cell.font = Font(italic=True)
    subtitle_cell.alignment = Alignment(horizontal='center')

    # Add headers
    headers = ["#", "Терапевт", "Ставка (%)", "Кол-во сеансов", "Оплачено (сум)", "К выплате (сум)"]
    for col_num, header in enumerate(headers, 1):
        cell = worksheet.cell(row=4, column=col_num)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")

    # Add data
    for row_num, therapist in enumerate(stats, 5):
        worksheet.cell(row=row_num, column=1).value = row_num - 4
        worksheet.cell(row=row_num, column=2).value = therapist['full_name']
        worksheet.cell(row=row_num, column=3).value = therapist['rate']
        worksheet.cell(row=row_num, column=4).value = therapist['session_count']
        worksheet.cell(row=row_num, column=5).value = therapist['total_amount']
        worksheet.cell(row=row_num, column=6).value = therapist['payout_amount']

    # Add totals
    total_row = len(stats) + 5
    worksheet.cell(row=total_row, column=1).value = "Итого:"
    worksheet.merge_cells(f'A{total_row}:C{total_row}')
    worksheet.cell(row=total_row, column=4).value = total_sessions
    worksheet.cell(row=total_row, column=5).value = total_amount
    worksheet.cell(row=total_row, column=6).value = total_payout

    # Make totals bold
    for col in range(1, 7):
        cell = worksheet.cell(row=total_row, column=col)
        cell.font = Font(bold=True)

    # Add profit distribution section
    dist_row = total_row + 2
    worksheet.cell(row=dist_row, column=1).value = "Распределение прибыли:"
    worksheet.merge_cells(f'A{dist_row}:F{dist_row}')
    worksheet.cell(row=dist_row, column=1).font = Font(bold=True)

    worksheet.cell(row=dist_row + 1, column=1).value = "Общая оплаченная сумма:"
    worksheet.merge_cells(f'A{dist_row + 1}:C{dist_row + 1}')
    worksheet.cell(row=dist_row + 1, column=5).value = total_amount

    worksheet.cell(row=dist_row + 2, column=1).value = "Доля партнера 1 (50%):"
    worksheet.merge_cells(f'A{dist_row + 2}:C{dist_row + 2}')
    worksheet.cell(row=dist_row + 2, column=5).value = owner1_share

    worksheet.cell(row=dist_row + 3, column=1).value = "Выплаты терапевтам:"
    worksheet.merge_cells(f'A{dist_row + 3}:C{dist_row + 3}')
    worksheet.cell(row=dist_row + 3, column=5).value = total_payout

    worksheet.cell(row=dist_row + 4, column=1).value = "Итоговая сумма партнера 1:"
    worksheet.merge_cells(f'A{dist_row + 4}:C{dist_row + 4}')
    worksheet.cell(row=dist_row + 4, column=5).value = owner1_final
    worksheet.cell(row=dist_row + 4, column=1).font = Font(bold=True)
    worksheet.cell(row=dist_row + 4, column=5).font = Font(bold=True)

    worksheet.cell(row=dist_row + 5, column=1).value = "Доля партнера 2 (50%):"
    worksheet.merge_cells(f'A{dist_row + 5}:C{dist_row + 5}')
    worksheet.cell(row=dist_row + 5, column=5).value = owner2_share
    worksheet.cell(row=dist_row + 5, column=1).font = Font(bold=True)
    worksheet.cell(row=dist_row + 5, column=5).font = Font(bold=True)

    # Create response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response[
        'Content-Disposition'] = f'attachment; filename=therapist_statistics_{start_date.strftime("%Y%m%d")}-{end_date.strftime("%Y%m%d")}.xlsx'

    # Save workbook to response
    workbook.save(response)

    return response


@login_required
def export_therapist_statistics_word(request):
    """
    Export therapist statistics to Word document with new profit distribution logic
    based on actual paid amounts
    """
    # Make sure python-docx is installed: pip install python-docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    # Default date range (last 30 days)
    end_date = datetime.now().date()
    start_date = end_date - timedelta(days=30)

    # Get date range from request if provided
    if 'date_range' in request.GET and request.GET['date_range']:
        date_range = request.GET['date_range']
        start_str, end_str = date_range.split(' - ')
        start_date = datetime.strptime(start_str, '%d/%m/%Y').date()
        end_date = datetime.strptime(end_str, '%d/%m/%Y').date()

    # Get all therapists and their statistics
    therapists = TherapistModel.objects.all()

    # Process statistics for each therapist
    stats = []
    total_sessions = 0
    total_amount = 0
    total_payout = 0

    for therapist in therapists:
        # Get sessions for this therapist in the date range
        sessions = SessionModel.objects.filter(
            therapist=therapist,
            created_at__date__gte=start_date,
            created_at__date__lte=end_date,
        )

        # Calculate total paid amount for these sessions
        session_ids = sessions.values_list('id', flat=True)
        paid_amount = PaymentModel.objects.filter(
            session_id__in=session_ids
        ).aggregate(total=Sum('amount'))['total'] or 0

        # Calculate statistics
        session_count = sessions.aggregate(count=Count('id'))['count'] or 0

        # Skip therapists with no sessions or no payments
        if session_count == 0 or paid_amount == 0:
            continue

        # Add to totals
        total_sessions += session_count
        total_amount += paid_amount

        # Calculate payout based on therapist rate (from one owner's share)
        rate = therapist.rate if hasattr(therapist, 'rate') else 0

        # New calculation: therapist gets their rate % from HALF of the paid amount
        owner_share = paid_amount / 2
        payout = owner_share * (rate / 100)
        total_payout += payout

        # Add to stats list
        stats.append({
            'full_name': therapist.full_name,
            'rate': rate,
            'session_count': session_count,
            'total_amount': paid_amount,
            'payout_amount': payout,
        })

    # Calculate split between owners
    owner1_share = total_amount / 2
    owner2_share = total_amount / 2

    # Owner 1 pays the therapists
    owner1_final = owner1_share - total_payout

    # Create a new Word document
    document = Document()

    # Add document title
    title = document.add_heading('Статистика терапевтов', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add date range info
    date_info = document.add_paragraph(f'Период: {start_date.strftime("%d/%m/%Y")} - {end_date.strftime("%d/%m/%Y")}')
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle about calculation method
    subtitle = document.add_paragraph('Расчет на основе фактически оплаченных сумм')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'

    # Add some space
    document.add_paragraph()

    # Add table headers
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = '#'
    header_cells[1].text = 'Терапевт'
    header_cells[2].text = 'Ставка (%)'
    header_cells[3].text = 'Кол-во сеансов'
    header_cells[4].text = 'Оплачено (сум)'
    header_cells[5].text = 'К выплате (сум)'

    # Make headers bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add data rows
    for idx, stat in enumerate(stats, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = stat['full_name']
        row_cells[2].text = f"{stat['rate']}%"
        row_cells[3].text = str(stat['session_count'])
        row_cells[4].text = f"{stat['total_amount']:,} сум".replace(',', ' ')
        row_cells[5].text = f"{stat['payout_amount']:,} сум".replace(',', ' ')

    # Add totals row
    totals_row = table.add_row().cells
    totals_row[0].merge(totals_row[2])
    totals_row[0].text = 'Итого:'
    totals_row[3].text = str(total_sessions)
    totals_row[4].text = f"{total_amount:,} сум".replace(',', ' ')
    totals_row[5].text = f"{total_payout:,} сум".replace(',', ' ')

    # Make totals bold
    for cell in totals_row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add profit distribution section
    document.add_paragraph()
    document.add_heading('Распределение прибыли', level=2)

    distribution = document.add_paragraph()
    distribution.add_run('Общая оплаченная сумма: ').bold = True
    distribution.add_run(f"{total_amount:,} сум\n".replace(',', ' '))

    distribution.add_run('Доля партнера 1 (50%): ').bold = True
    distribution.add_run(f"{owner1_share:,} сум\n".replace(',', ' '))

    distribution.add_run('Выплаты терапевтам: ').bold = True
    distribution.add_run(f"{total_payout:,} сум\n".replace(',', ' '))

    distribution.add_run('Итоговая сумма партнера 1: ').bold = True
    distribution.add_run(f"{owner1_final:,} сум\n".replace(',', ' '))

    distribution.add_run('Доля партнера 2 (50%): ').bold = True
    distribution.add_run(f"{owner2_share:,} сум".replace(',', ' '))

    # Add footer
    document.add_paragraph()
    footer = document.add_paragraph()
    footer.add_run(f"Отчет сгенерирован: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Create response with Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response[
        'Content-Disposition'] = f'attachment; filename=therapist_statistics_{start_date.strftime("%Y%m%d")}-{end_date.strftime("%Y%m%d")}.docx'

    # Save document to response
    document.save(response)

    return response